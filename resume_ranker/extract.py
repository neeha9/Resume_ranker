"""Extract plain text from uploaded resume files (PDF, DOCX, TXT)."""

from __future__ import annotations

import io

import docx
from pypdf import PdfReader


def extract_text(filename: str, data: bytes) -> str:
    """Return the plain text content of a resume file.

    Raises ValueError for unsupported or unreadable files so callers can
    surface a per-file error instead of failing the whole batch.
    """
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if suffix == "pdf":
        return _extract_pdf(data)
    if suffix == "docx":
        return _extract_docx(data)
    if suffix == "txt":
        return data.decode("utf-8", errors="replace")

    raise ValueError(f"Unsupported file type: .{suffix}")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages).strip()
    if not text:
        raise ValueError("No extractable text found (the PDF may be a scanned image).")
    return text


def _extract_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(p for p in parts if p.strip())
    if not text.strip():
        raise ValueError("No extractable text found in the document.")
    return text
